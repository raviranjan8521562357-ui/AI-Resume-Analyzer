"""
Resume text extraction from PDF and DOCX files.
Enhanced with better error handling and debug logging.
"""
import io
from pypdf import PdfReader
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file with robust error handling."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        print(f"[extractor] Failed to read PDF: {e}")
        raise ValueError(f"Could not read PDF file: {e}")

    text = ""
    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text()
            if page_text:
                # Clean up common PDF extraction artifacts
                # Fix words that got split across lines without spaces
                page_text = page_text.replace('\x00', '')  # Remove null bytes
                text += page_text + "\n"
        except Exception as e:
            print(f"[extractor] Warning: Failed to extract page {i+1}: {e}")
            continue

    # Debug: log extraction quality
    line_count = len([l for l in text.split('\n') if l.strip()])
    char_count = len(text.strip())
    print(f"[extractor] PDF extracted: {char_count} chars, {line_count} non-empty lines")

    if char_count < 50:
        print(f"[extractor] WARNING: Very little text extracted ({char_count} chars). "
              f"PDF may be image-based or have unusual encoding.")

    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        print(f"[extractor] Failed to read DOCX: {e}")
        raise ValueError(f"Could not read DOCX file: {e}")

    text = ""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    # Also extract text from tables (resumes often use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text += " | ".join(row_text) + "\n"

    char_count = len(text.strip())
    print(f"[extractor] DOCX extracted: {char_count} chars")

    if char_count < 50:
        print(f"[extractor] WARNING: Very little text from DOCX ({char_count} chars).")

    return text


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Extract text from a file based on its extension.
    
    Args:
        filename: Original filename (used to detect format)
        file_bytes: Raw file bytes
    
    Returns:
        Extracted text string
    
    Raises:
        ValueError: If file format is not supported
    """
    filename_lower = filename.lower()
    print(f"[extractor] Processing: {filename} ({len(file_bytes)} bytes)")
    
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format. Please upload a PDF or DOCX file.")
